import openai
from gateway.config import *
import docx
from fastapi import FastAPI, UploadFile, File, Depends, FastAPI, HTTPException, APIRouter
from typing import Optional
from fastapi.security.http import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel
from pypdf import PdfReader

api_app = FastAPI()

auth_token_handler = HTTPBearer(auto_error=False)

async def get_token(auth: Optional[HTTPAuthorizationCredentials] = Depends(auth_token_handler)) -> str:
    if auth is None or auth.credentials != AUTH_SWAGGER_TOKEN:
        raise HTTPException(401, detail="Bearer token missing or unknown")
    return auth.credentials


docu_flow_router = APIRouter(
    tags=["DocuFlow"],
    dependencies=[Depends(get_token)]
)



class extract_data_from_file_result(BaseModel):
    document_text: str

class get_explanation_result(BaseModel):
    explanation_result: str


@docu_flow_router.post("/extract_data_from_file/")
async def extract_data_from_file(input_file: UploadFile = File(...)) -> extract_data_from_file_result:
    if input_file.content_type == "application/pdf":
        document_text =  '\n'.join(pdf_page.extract_text() for pdf_page in PdfReader(input_file.file).pages)
    elif input_file.content_type == "text/plain":
        document_text = input_file.file.read().decode('utf-8')
    elif input_file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document_text =  '\n'.join(paragraph.text for paragraph in docx.Document(input_file.file).paragraphs)
    else:
        raise HTTPException(400, detail="Invalid document type")
    return {"document_text": document_text}



@docu_flow_router.get("/get_explanation")
async def get_explanation(document_text: str) -> get_explanation_result:
    try:
        chat_gpt_request = openai.OpenAI(api_key=CHAT_GPT_TOKEN)

        explanation_result = chat_gpt_request.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "assistant", "content": 'document:{document_text}.(Country Ukraine).\nMake big explanation of this document in simple words for an ordinary person, but do not forget about the details (names, numbers, etc.). As a result, you will receive a text about this document in the language of the country.'}
            ]
        ).choices[0].message.content
    except Exception as explanation_error:
        raise HTTPException(400, detail=f"Explanation was failed\n\n{explanation_error}")

    return {"explanation_result": explanation_result}




api_app.include_router(docu_flow_router)

# document_text =  '\n'.join(paragraph.text for paragraph in docx.Document('/Users/dmitro/Downloads/Договір з Ліцензіатом.docx').paragraphs)

# chat_gpt_request = openai.OpenAI(api_key=CHAT_GPT_TOKEN)


# completion = chat_gpt_request.chat.completions.create(
#     model="gpt-4o",
#     messages=[
#         {"role": "system", "content":document_text},
#         {"role": "assistant", "content": '(Country Ukraine) Explain this document in simple words for an ordinary person, but do not forget about the details (names, numbers, etc.). As a result, you will receive a text about this document in the language of the country.'}
#     ]
# ).choices[0].message.content

# print(completion)